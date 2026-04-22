import streamlit as st
import xml.etree.ElementTree as ET
from docx import Document
from docx.document import Document as DocxDocument
from docx.oxml.table import CT_Tbl
from docx.oxml.text.paragraph import CT_P
from docx.table import _Cell, Table
from docx.text.paragraph import Paragraph
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import io
import json
import os
from collections import deque, defaultdict
import tempfile
import re

def _iter_block_items(parent):
    """Walk body in document order (paragraphs and tables interleaved)."""
    if isinstance(parent, DocxDocument):
        body = parent.element.body
    elif isinstance(parent, _Cell):
        body = parent._tc
    else:
        return
    for child in body.iterchildren():
        if isinstance(child, CT_P):
            yield ("p", Paragraph(child, parent))
        elif isinstance(child, CT_Tbl):
            yield ("t", Table(child, parent))


_RE_DOCX_Q = re.compile(
    r"^([A-Za-z]+\d+[a-zA-Z]*)\s*[:.]\s*(.+)$",
    re.DOTALL,
)
_META_LINE_PREFIXES = (
    "Q Type:",
    "Rotate/Randomize:",
    "Programmer notes:",
    "Тип запитання:",
    "Ротація/Рандомізація:",
    "Примітки програміста:",
)


def _docx_is_meta_line(txt):
    t = txt.strip()
    return any(t.startswith(p) for p in _META_LINE_PREFIXES)


# Short standalone caps lines (mode labels) must not become extra "sections" or they
# desync EN vs UA and shift every following row in the bilingual table.
_SECTION_SINGLE_WORD_DENY = frozenset({
    "OFFLINE", "ONLINE", "ОФЛАЙН", "ОНЛАЙН", "CATI", "CAWI", "CAPI",
})


def _docx_looks_like_section(txt):
    t = txt.strip()
    if _docx_is_meta_line(t) or _RE_DOCX_Q.match(t):
        return False
    if len(t) < 4 or len(t) > 140:
        return False
    if t.upper() in _SECTION_SINGLE_WORD_DENY:
        return False
    if re.match(r"^Part\s+[A-Z]\.", t):
        return True
    letters = [c for c in t if c.isalpha()]
    if len(letters) < 4:
        return False
    upper_ratio = sum(1 for c in letters if c.isupper()) / len(letters)
    return upper_ratio >= 0.72


def _docx_looks_like_option_paragraph(line, stem_first_line):
    """
    Heuristic: lines after the script question that are answer rows / show lines / inputs.
    Avoid treating a single long continuation paragraph of the stem as an option.
    """
    s = (line or "").strip()
    if not s:
        return False
    if _docx_is_meta_line(s):
        return False
    # Numbered / lettered options: 1. 2) 99. a)
    if re.match(r"^(\d+|[a-zA-Z])\s*[\).:]\s+\S", s):
        return True
    if s.startswith("["):
        return True
    # Script markers
    if re.search(r"\[EXCLUSIVE\]|\[exclusive\]", s, re.I):
        return True
    stem = (stem_first_line or "").strip()
    stem_is_question = bool(stem) and stem.rstrip().endswith("?")
    # After a question stem, follow-on lines are usually options / prompts (allow long statements)
    if stem_is_question and len(s) <= 900 and not s.endswith("?"):
        if len(s) > 360:
            return bool(
                re.match(r"^(\d+|[a-zA-Z])\s*[\).:]\s", s)
                or s.startswith("[")
                or re.match(
                    r"^(I am|We |This is|None of|The |Other|Air |Ocean |Land )",
                    s,
                    re.I,
                )
                or re.match(r"^(Я |Ми |Це |Інш|Жодн)", s, re.I)
            )
        return True
    # Statement-style script (stem ends with .) — still often followed by fixed statements
    if stem.rstrip().endswith(".") and len(s) <= 900:
        if re.match(
            r"^(I am|We |This is|None of|The |Other|Air |Ocean |Land |\d|\[|Я |Ми |Це |Інш|Жодн)",
            s,
            re.I,
        ):
            return True
    # Percent / open input lines
    if re.search(r"_%\]|____%|___%", s):
        return True
    return False


def _docx_table_to_answers(qid, table):
    answers = []
    for row in table.rows:
        if not row.cells:
            continue
        code = row.cells[0].text.strip()
        label = (
            row.cells[1].text.replace("\n", " ").strip()
            if len(row.cells) > 1
            else ""
        )
        if not code and not label:
            continue
        if re.match(r"^[\d\-]+$", code):
            answers.append({
                "id": f"{qid}_{code}",
                "role": "row",
                "number": code,
                "text": label,
            })
    return answers


def _docx_flush_question_buffer(current_q, buf):
    """Turn queued paragraphs/tables after a question header into stem text + answers."""
    if not current_q or not buf:
        return

    while buf and buf[0][0] == "p" and _docx_is_meta_line(buf[0][1]):
        buf.pop(0)
    while buf and buf[-1][0] == "p" and _docx_is_meta_line(buf[-1][1]):
        buf.pop()
    if not buf:
        return

    has_table = any(k == "t" for k, _ in buf)

    if has_table:
        stem_parts = []
        for k, d in buf:
            if k == "p" and not _docx_is_meta_line(d):
                stem_parts.append(d.strip())
            elif k == "t":
                break
        current_q["text"] = "\n".join(stem_parts).strip()
        for k, d in buf:
            if k == "t":
                current_q["answers"].extend(_docx_table_to_answers(current_q["id"], d))
        return

    split_at = None
    for i, (k, t) in enumerate(buf):
        if k == "p" and _docx_is_meta_line(t):
            split_at = i
            break

    texts = [d.strip() for k, d in buf if k == "p" and not _docx_is_meta_line(d)]

    if split_at is None:
        if not texts:
            return
        if len(texts) == 1:
            current_q["text"] = texts[0]
            return
        stem_first = texts[0]
        i = 1
        stem_parts = [stem_first]
        while i < len(texts) and not _docx_looks_like_option_paragraph(texts[i], stem_first):
            stem_parts.append(texts[i])
            i += 1
        current_q["text"] = "\n".join(stem_parts).strip()
        opt_lines = texts[i:]
        for j, ot in enumerate(opt_lines, 1):
            current_q["answers"].append({
                "id": f"{current_q['id']}_{j}",
                "role": "row",
                "number": str(j),
                "text": ot,
            })
        return

    pre_texts = [
        d.strip()
        for k, d in buf[:split_at]
        if k == "p" and not _docx_is_meta_line(d)
    ]
    if not pre_texts:
        return
    current_q["text"] = pre_texts[0]
    for j, ot in enumerate(pre_texts[1:], 1):
        current_q["answers"].append({
            "id": f"{current_q['id']}_{j}",
            "role": "row",
            "number": str(j),
            "text": ot,
        })


def parse_docx(file_bytes):
    """
    Parse script-style questionnaires: question lines like ``S1: …`` or ``D4a. …``,
    tables with coded rows, and paragraph option lists (EN/UA use the same question IDs).
    """
    tmp_path = None
    try:
        with tempfile.NamedTemporaryFile(delete=False, suffix=".docx") as tmp:
            tmp.write(file_bytes)
            tmp_path = tmp.name

        doc = Document(tmp_path)

        items = []
        sec_idx = 0
        section_id = "SEC_0"
        current_q = None
        buf = []

        for kind, item in _iter_block_items(doc):
            if kind == "p":
                t = item.text.strip()
                if not t:
                    continue

                if _docx_looks_like_section(t):
                    _docx_flush_question_buffer(current_q, buf)
                    buf = []
                    current_q = None
                    sid = f"SEC_{sec_idx}"
                    sec_idx += 1
                    section_id = sid
                    items.append({
                        "id": sid,
                        "type": "section",
                        "sectionId": sid,
                        "text": t,
                        "answers": [],
                    })
                    continue

                m = _RE_DOCX_Q.match(t)
                if m:
                    _docx_flush_question_buffer(current_q, buf)
                    buf = []
                    qid = m.group(1)
                    title = m.group(2).strip()
                    current_q = {
                        "id": qid,
                        "type": "question",
                        "sectionId": section_id,
                        "text": title,
                        "answers": [],
                    }
                    items.append(current_q)
                    continue

                if current_q:
                    buf.append(("p", t))
            else:
                if current_q:
                    buf.append(("t", item))

        _docx_flush_question_buffer(current_q, buf)

        return items
    finally:
        if tmp_path and os.path.isfile(tmp_path):
            try:
                os.unlink(tmp_path)
            except OSError:
                pass

def parse_file(uploaded_file):
    ext = uploaded_file.name.lower().split('.')[-1]

    if ext == "xml":
        return parse_xml(uploaded_file.read())

    if ext == "docx":
        return parse_docx(uploaded_file.read())

    raise Exception("Unsupported file type")
st.set_page_config(page_title="Survey Bilingual Generator", layout="centered")
st.title("📋 Survey Bilingual Document Generator")
st.markdown(
    "Завантажте два файли опитування (**XML** або **DOCX**): перший — основна мова, "
    "другий — переклад. Для Word очікується структура як у скрипті (рядки **S1:**, **D4a.** тощо, "
    "таблиці з кодами у першій колонці або списки варіантів до наступного питання)."
)

# ── HELPERS ──────────────────────────────────────────────────────────────────

NON_QUESTION_KEYS = {
    'id', 'title', 'current_version_id', 'current_version_number',
    'next_version_id', 'num_of_versions', 'plotly_export_format',
    'previous_version_id', 'regenerate_collection', 'section_type', 'methodology'
}

def extract_title(elem):
    """Extract question text from <title> or <textblockcontent>."""
    title = elem.find('title')
    if title is not None:
        p = title.find('paragraph')
        if p is not None:
            return get_full_text(p)
    tbc = elem.find('textblockcontent')
    if tbc is not None:
        p = tbc.find('paragraph')
        if p is not None:
            return get_full_text(p)
    return ''

def get_full_text(elem):
    """Recursively get all text from element."""
    parts = []
    if elem.text:
        parts.append(elem.text.strip())
    for child in elem:
        parts.append(get_full_text(child))
        if child.tail:
            parts.append(child.tail.strip())
    return ' '.join(p for p in parts if p)

def extract_answers(elem):
    """Extract all answer rows from any answers structure."""
    answers = []
    answers_node = elem.find('answers')
    if answers_node is None:
        # textquestion uses <row> directly
        for row in elem.findall('row'):
            answers.append({
                'id': row.get('id', ''),
                'role': 'row',
                'text': get_full_text(row)
            })
        return answers

    # answers can be repeated (grid has multiple answers blocks)
    all_answer_nodes = elem.findall('answers')
    for ans_node in all_answer_nodes:
        child_type = ans_node.get('childrenType', 'row')
        for child_tag in ['row', 'grow', 'gcol']:
            for item in ans_node.findall(child_tag):
                answers.append({
                    'id': item.get('id', ''),
                    'role': child_type,
                    'text': get_full_text(item)
                })
    return answers

def parse_xml(xml_bytes):
    """Parse XML bytes into normalized list of items."""
    root = ET.fromstring(xml_bytes)
    items = []

    for section in root.findall('.//section'):
        section_id = section.get('id', '')
        section_title = section.get('title', section_id)

        items.append({
            'id': section_id,
            'type': 'section',
            'sectionId': section_id,
            'text': section_title,
            'answers': []
        })

        # Dynamically find all question-type children
        for child in section:
            tag = child.tag
            if tag in NON_QUESTION_KEYS:
                continue
            q_id = child.get('id')
            if not q_id:
                continue

            items.append({
                'id': q_id,
                'type': tag,
                'sectionId': section_id,
                'text': extract_title(child),
                'answers': extract_answers(child)
            })

    return items

def _answer_code_from_primary_row(ans):
    """Numeric code from table column or from id suffix (e.g. S1_37 -> 37)."""
    num = (ans.get('number') or '').strip()
    if num and re.match(r'^[\d\-]+$', num):
        return num
    aid = ans.get('id', '')
    if '_' in aid:
        suf = aid.rsplit('_', 1)[-1]
        if re.match(r'^[\d\-]+$', suf):
            return suf
    return ''

def _leading_numbered_code_in_text(text):
    """UA sometimes keeps script code as ``37. text`` — align to EN row code 37."""
    m = re.match(r'^\s*(\d+)\s*[).:]\s+', text or '')
    if m:
        return m.group(1)
    return None

def _collapse_bracket_option_runs(texts):
    """
    Merge consecutive short options that start with '[' (split Show/Показувати lines)
    so two EN table rows map to two combined UA blocks instead of four UA lines.
    """
    out = []
    i = 0
    while i < len(texts):
        t = (texts[i] or '').strip()
        if i + 1 < len(texts):
            u = (texts[i + 1] or '').strip()
            if t.startswith('[') and u.startswith('['):
                out.append(texts[i].rstrip() + '\n' + texts[i + 1].rstrip())
                i += 2
                continue
        out.append(texts[i])
        i += 1
    return out

def _merge_answer_lists_aligned(primary_answers, secondary_answers):
    """
    Pair EN/UA answer rows for DOCX where ids are ordinal on one side and coded on the other.
    Uses: (1) explicit numbers at start of UA line -> EN code; (2) merge consecutive '[' lines;
    (3) FIFO for remaining rows in document order.
    """
    if not primary_answers:
        return []
    if not secondary_answers:
        return [
            {
                'id': a['id'],
                'role': a['role'],
                'number': a.get('number', ''),
                'primary': a['text'],
                'secondary': '',
            }
            for a in primary_answers
        ]

    sec_texts = [a.get('text', '') for a in secondary_answers]
    keyed = {}
    non_keyed = []
    for t in sec_texts:
        code = _leading_numbered_code_in_text(t)
        if code is not None:
            keyed[code] = t.strip()
        else:
            non_keyed.append(t)
    non_keyed = _collapse_bracket_option_runs(non_keyed)
    plain = deque(non_keyed)
    used_keyed = set()

    rows = []
    for pa in primary_answers:
        code = _answer_code_from_primary_row(pa)
        sec = ''
        if code and code in keyed and code not in used_keyed:
            sec = keyed[code]
            used_keyed.add(code)
        elif plain:
            sec = plain.popleft().strip()
        rows.append({
            'id': pa['id'],
            'role': pa.get('role', 'row'),
            'number': pa.get('number', ''),
            'primary': pa['text'],
            'secondary': sec,
        })
    return rows

def _answers_merge_by_ids_ok(primary_answers, secondary_answers):
    """True when ids line up (same count + almost every id has non-empty secondary text)."""
    if not primary_answers:
        return True
    if len(primary_answers) != len(secondary_answers):
        return False
    sm = {a['id']: a.get('text', '') for a in secondary_answers}
    hits = sum(1 for a in primary_answers if sm.get(a['id'], '').strip())
    return hits / len(primary_answers) >= 0.92

def _split_secondary_question_body_to_answers(s_item, pa):
    """
    If options were merged into question ``text`` (legacy flush / Word layout),
    move trailing lines into answer rows aligned with ``pa`` by count/order.
    """
    if not s_item or not pa:
        return s_item, []
    sa = list(s_item.get('answers') or [])
    if len(sa) >= len(pa):
        return s_item, sa
    body = (s_item.get('text') or '').strip()
    if not body:
        return s_item, sa
    lines = [ln.strip() for ln in body.splitlines() if ln.strip()]
    if len(lines) < 2 or len(lines) - 1 < len(pa):
        return s_item, sa
    head, tail = lines[0], lines[1:]
    if len(tail) < len(pa):
        return s_item, sa
    fixed = dict(s_item)
    fixed['text'] = head
    new_sa = []
    for i, a in enumerate(pa):
        new_sa.append({
            'id': a['id'],
            'role': a.get('role', 'row'),
            'number': a.get('number', ''),
            'text': tail[i] if i < len(tail) else '',
        })
    return fixed, new_sa


def merge(primary_items, secondary_items):
    """Merge primary and secondary by ID (FIFO when the same id appears more than once)."""
    secondary_queues = defaultdict(deque)
    for item in secondary_items:
        secondary_queues[item['id']].append(item)

    merged = []
    for p_item in primary_items:
        qid = p_item['id']
        qsec = secondary_queues.get(qid)
        s_item = qsec.popleft() if qsec else None

        pa = p_item.get('answers', [])
        sa = list(s_item.get('answers', [])) if s_item else []

        if s_item is not None:
            s_item, sa = _split_secondary_question_body_to_answers(s_item, pa)

        secondary_answer_map = {a['id']: (a.get('text') or '') for a in sa}

        if pa and sa and not _answers_merge_by_ids_ok(pa, sa):
            merged_ans = _merge_answer_lists_aligned(pa, sa)
        else:
            merged_ans = [
                {
                    'id': a['id'],
                    'role': a['role'],
                    'number': a.get('number', ''),
                    'primary': a['text'],
                    'secondary': secondary_answer_map.get(a['id'], ''),
                }
                for a in pa
            ]

        merged.append({
            'id': p_item['id'],
            'type': p_item['type'],
            'sectionId': p_item['sectionId'],
            'primary': p_item['text'],
            'secondary': s_item['text'] if s_item else '',
            'answers': merged_ans,
        })
    return merged

def set_cell_bg(cell, hex_color):
    """Set cell background color."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)

def set_col_width(table, col_index, width_inches):
    """Set column width."""
    for row in table.rows:
        row.cells[col_index].width = Inches(width_inches)


def _format_question_heading(qid, body_text):
    """Same style as reference: ``S4. Question text`` in both language columns."""
    qid = (qid or "").strip()
    body = (body_text or "").strip()
    if not qid:
        return body
    if not body:
        return f"{qid}."
    return f"{qid}. {body}"


def _set_table_width_percent(table, pct_fiftieths):
    """pct_fiftieths: 5000 = 100% of parent (Word ``pct`` width)."""
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    if tbl_pr is None:
        tbl_pr = OxmlElement("w:tblPr")
        tbl.insert(0, tbl_pr)
    for child in list(tbl_pr):
        if child.tag == qn("w:tblW"):
            tbl_pr.remove(child)
    tbl_w = OxmlElement("w:tblW")
    tbl_w.set(qn("w:type"), "pct")
    tbl_w.set(qn("w:w"), str(pct_fiftieths))
    tbl_pr.append(tbl_w)


def _add_nested_answer_table(cell, answers, use_primary_text, label_column_width=None):
    """
    Nested 2-column table: index | label.
    Rows follow the primary (English) answer list so left/right stay aligned.
    Table is ~full width of the cell and centered.
    """
    leaders = [a for a in answers if (a.get("primary") or "").strip()]
    if not leaders:
        cell.text = ""
        return
    nt = cell.add_table(rows=len(leaders), cols=2)
    nt.style = "Table Grid"
    nt.autofit = False
    nt.alignment = WD_TABLE_ALIGNMENT.CENTER
    _set_table_width_percent(nt, 4800)

    idx_w = Inches(0.52)
    if label_column_width is not None:
        lbl_w = Inches(max(1.8, label_column_width.inches - idx_w.inches))
    else:
        lbl_w = Inches(3.0)

    for ri, ans in enumerate(leaders):
        num = (ans.get("number") or "").strip() or str(ri + 1)
        if use_primary_text:
            label = (ans.get("primary") or "").strip()
        else:
            label = (ans.get("secondary") or "").strip()
        row = nt.rows[ri]
        row.cells[0].text = num
        row.cells[1].text = label
        for c in row.cells:
            for para in c.paragraphs:
                for run in para.runs:
                    run.font.size = Pt(9)
        row.cells[0].width = idx_w
        row.cells[1].width = lbl_w


def generate_docx(merged, primary_label, secondary_label):
    """Generate bilingual DOCX from merged data."""
    doc = Document()

    # Page: landscape (wider columns for two languages)
    section = doc.sections[0]
    section.left_margin = Inches(0.7)
    section.right_margin = Inches(0.7)
    section.top_margin = Inches(0.7)
    section.bottom_margin = Inches(0.7)
    new_width, new_height = section.page_height, section.page_width
    section.orientation = WD_ORIENT.LANDSCAPE
    section.page_width = new_width
    section.page_height = new_height

    col_half = Inches(
        (section.page_width.inches - section.left_margin.inches - section.right_margin.inches)
        / 2.0
    )

    # Title
    title = doc.add_heading('Bilingual Survey Document', level=1)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph('')

    # Table
    table = doc.add_table(rows=1, cols=2)
    table.style = 'Table Grid'

    # Header row
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = primary_label
    hdr_cells[1].text = secondary_label

    for cell in hdr_cells:
        set_cell_bg(cell, '2C3E50')
        for para in cell.paragraphs:
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in para.runs:
                run.bold = True
                run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                run.font.size = Pt(11)

    # Data rows
    for item in merged:
        if item['type'] == 'section':
            row = table.add_row()
            cells = row.cells
            sec = (item.get('secondary') or '').strip()
            if sec:
                cells[0].text = f"  {item['primary'].upper()}"
                cells[1].text = f"  {sec.upper()}"
                for cell in cells:
                    set_cell_bg(cell, '1ABC9C')
                    for para in cell.paragraphs:
                        for run in para.runs:
                            run.bold = True
                            run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                            run.font.size = Pt(10)
            else:
                cells[0].merge(cells[1])
                cells[0].text = f"  {item['primary'].upper()}"
                set_cell_bg(cells[0], '1ABC9C')
                for para in cells[0].paragraphs:
                    for run in para.runs:
                        run.bold = True
                        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                        run.font.size = Pt(10)
            continue

        if item['type'] == 'textblock':
            if not item['primary']:
                continue
            row = table.add_row()
            cells = row.cells
            cells[0].text = item['primary']
            cells[1].text = item['secondary']
            set_cell_bg(cells[0], 'ECF0F1')
            set_cell_bg(cells[1], 'ECF0F1')
            for cell in cells:
                for para in cell.paragraphs:
                    for run in para.runs:
                        run.italic = True
                        run.font.size = Pt(9)
            continue

        # Questions: heading row, then next row = nested option tables (no spacer row)
        qid = str(item.get("id", ""))
        answers = item.get("answers") or []
        prim = (item.get("primary") or "").strip()
        sec = (item.get("secondary") or "").strip()
        has_answers = any((a.get("primary") or "").strip() for a in answers)

        if not prim and not has_answers:
            continue

        qrow = table.add_row()
        qc = qrow.cells
        qc[0].text = _format_question_heading(qid, prim) if prim else f"{qid}."
        qc[1].text = _format_question_heading(qid, sec) if sec else f"{qid}."
        for cell in qc:
            set_cell_bg(cell, "EBF5FB")
            for para in cell.paragraphs:
                for run in para.runs:
                    run.bold = True
                    run.font.size = Pt(10)

        if has_answers:
            opt_row = table.add_row()
            oc0, oc1 = opt_row.cells[0], opt_row.cells[1]
            for c in (oc0, oc1):
                set_cell_bg(c, "FFFFFF")
            _add_nested_answer_table(oc0, answers, True, col_half)
            _add_nested_answer_table(oc1, answers, False, col_half)

    for row in table.rows:
        for cell in row.cells:
            cell.width = col_half

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf

# ── UI ─────────────────────────────────────────

col1, col2 = st.columns(2)

with col1:
    primary_label = st.text_input(
        "Primary language label",
        value="English"
    )

    primary_file = st.file_uploader(
        "Upload Primary file",
        type=['xml', 'docx']
    )

with col2:
    secondary_label = st.text_input(
        "Secondary language label",
        value="Translation"
    )

    secondary_file = st.file_uploader(
        "Upload Secondary file",
        type=['xml', 'docx']
    )

if primary_file and secondary_file:

    if st.button("🚀 Generate Bilingual DOCX", type="primary"):

        with st.spinner("Parsing and merging..."):

            try:
                primary_items = parse_file(primary_file)
                secondary_items = parse_file(secondary_file)

                merged = merge(primary_items, secondary_items)

                st.success(
                    f"✅ Merged {len(merged)} items successfully"
                )

                if len(merged) == 0:
                    st.warning(
                        "Не вдалося розпізнати структуру DOCX/XML. "
                        "Для Word: питання як **S1: …** або **D4a. …**, секції — переважно ВЕЛИКИМИ ЛІТЕРАМИ, "
                        "варіанти — абзаци до «Q Type» / «Rotate» або рядки таблиці «код | текст». "
                        "Спробуйте експорт у XML, якщо макет інший."
                    )

                with st.expander("Preview merged data (first 10 items)"):
                    st.json(merged[:10])

                with st.spinner("Generating DOCX..."):
                    docx_buf = generate_docx(
                        merged,
                        primary_label,
                        secondary_label
                    )

                if docx_buf is not None:
                    st.download_button(
                        label="📥 Download DOCX",
                        data=docx_buf,
                        file_name="survey_bilingual.docx",
                        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                    )

            except Exception as e:
                st.error(f"Error: {str(e)}")
                st.exception(e)

else:
    st.info("👆 Upload both files to get started.")
